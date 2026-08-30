# ruff: noqa: E402, I001
from __future__ import annotations

import base64
from dataclasses import dataclass
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

pytest.importorskip("PySide6")

from PySide6.QtCore import Qt
from PySide6.QtTest import QTest

from page_line_editor.pagexml.parser import PAGE_2013_NAMESPACE, parse_page
from page_line_editor.ui.controller import (
    EditorController,
    _BatchProposal,
    _BatchResult,
    _source_digest,
)
from page_line_editor.ui.main_window import MainWindow
from page_line_editor.ui.panels import ProjectPaths


# A real, synthetic 1x1 PNG. PAGE dimensions deliberately remain document
# coordinates; the controller should not need manuscript fixtures to open it.
PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _page_xml(text: str = "قديم") -> bytes:
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<PcGts xmlns="{PAGE_2013_NAMESPACE}">
  <Metadata>
    <Creator>controller test</Creator>
    <Created>2020-01-01T00:00:00Z</Created>
    <LastChange>2020-01-01T00:00:00Z</LastChange>
  </Metadata>
  <Page imageFilename="1r.png" imageWidth="100" imageHeight="80">
    <TextRegion id="region-1">
      <Coords points="0,0 99,0 99,79 0,79"/>
      <TextLine id="line-1">
        <Coords points="10,10 90,10 90,35 10,35"/>
        <Baseline points="12,30 88,30"/>
        <TextEquiv><Unicode>{text}</Unicode></TextEquiv>
      </TextLine>
    </TextRegion>
  </Page>
</PcGts>'''.encode()


@dataclass(frozen=True)
class SyntheticProject:
    paths: ProjectPaths
    xml_path: Path
    original_xml: bytes


def _project(tmp_path: Path) -> SyntheticProject:
    images = tmp_path / "images"
    xml = tmp_path / "xml"
    history = tmp_path / "history"
    images.mkdir()
    xml.mkdir()
    (images / "1r.png").write_bytes(PNG_BYTES)
    original = _page_xml()
    xml_path = xml / "1r.xml"
    xml_path.write_bytes(original)

    ground_truth = tmp_path / "ground-truth.docx"
    document = Document()
    document.add_paragraph("[1r]")
    document.add_paragraph("جديد")
    document.save(ground_truth)
    return SyntheticProject(
        ProjectPaths(images, xml, ground_truth, history, normalize_nfc=True),
        xml_path,
        original,
    )


def _open_controller(qtbot, project: SyntheticProject) -> EditorController:  # type: ignore[no-untyped-def]
    window = MainWindow()
    qtbot.addWidget(window)
    controller = EditorController(window)
    controller.open_project(project.paths)
    assert controller.session.document is not None
    assert controller.ground_truth is not None
    return controller


def _complete_current_page_correction(controller: EditorController) -> None:
    """Run the worker's pure step and its GUI-thread completion deterministically."""
    document = controller.session.document
    assert document is not None
    assert controller.ground_truth is not None
    proposal = controller.workflow.propose(document, controller.ground_truth)
    controller._apply_page_proposal(proposal)


def _saved_unicode(path: Path) -> str:
    tree = etree.parse(str(path))
    return str(
        tree.xpath(
            "string(.//p:TextLine[@id='line-1']/p:TextEquiv/p:Unicode)",
            namespaces={"p": PAGE_2013_NAMESPACE},
        )
    )


def test_controller_opens_folders_auto_applies_in_memory_and_rejects(qtbot, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    project = _project(tmp_path)
    controller = _open_controller(qtbot, project)

    document = controller.session.document
    assert document is not None
    assert document.line_by_id("line-1").text == "قديم"
    _complete_current_page_correction(controller)

    line = document.line_by_id("line-1")
    assert line.text == "جديد"
    assert line.proposal_state == "applied"
    assert line.diff_text == "قديم → جديد"
    assert line.correction_status == "OCR"
    assert document.is_dirty
    assert project.xml_path.read_bytes() == project.original_xml
    # Auto correction immediately keeps its own audit copy/report, while the
    # live PAGE source remains untouched until explicit Save.
    assert (
        next(project.paths.audit_directory.glob("auto/*/originals/1r.xml")).read_bytes()
        == project.original_xml
    )

    controller.reject_line("line-1")
    assert line.text == "قديم"
    assert line.proposal_state == ""
    assert line.correction_status == ""
    assert not document.is_dirty
    assert project.xml_path.read_bytes() == project.original_xml
    assert not list(project.paths.audit_directory.glob("manual/*/originals/1r.xml"))

    controller.window.undo_stack.undo()
    assert line.text == "جديد"
    controller.window.undo_stack.redo()
    assert line.text == "قديم"


def test_controller_explicit_save_writes_correction_and_exact_manual_backup(
    qtbot, tmp_path: Path
) -> None:  # type: ignore[no-untyped-def]
    project = _project(tmp_path)
    controller = _open_controller(qtbot, project)
    _complete_current_page_correction(controller)

    assert project.xml_path.read_bytes() == project.original_xml
    assert not list(project.paths.audit_directory.glob("manual/*/originals/1r.xml"))
    controller.save()

    assert _saved_unicode(project.xml_path) == "جديد"
    backups = list(project.paths.audit_directory.glob("manual/*/originals/1r.xml"))
    assert len(backups) == 1
    assert backups[0].read_bytes() == project.original_xml
    assert controller.session.document is not None
    assert not controller.session.document.is_dirty
    assert controller.window.undo_stack.isClean()
    assert controller.window.undo_stack.count() == 0


def test_accepting_pending_extra_removes_geometry_and_saved_xml(
    qtbot, tmp_path: Path
) -> None:  # type: ignore[no-untyped-def]
    project = _project(tmp_path)
    ground_truth = Document()
    ground_truth.add_paragraph("[1r]")
    assert project.paths.ground_truth_path is not None
    ground_truth.save(project.paths.ground_truth_path)
    controller = _open_controller(qtbot, project)
    _complete_current_page_correction(controller)
    document = controller.session.document
    assert document is not None
    line = document.line_by_id("line-1")

    assert line.proposal_state == "pending"
    assert not line.deleted
    item = controller.window.canvas.page_scene.line_item("line-1")
    assert item is not None
    item.setSelected(True)
    controller.window.show()
    controller.window.activateWindow()
    controller.window.canvas.viewport().setFocus()
    qtbot.waitUntil(controller.window.canvas.viewport().hasFocus)

    QTest.keyClick(controller.window.canvas.viewport(), Qt.Key.Key_Return)

    assert line.deleted
    assert controller.window.canvas.page_scene.line_item("line-1") is None
    assert project.xml_path.read_bytes() == project.original_xml
    controller.save()
    tree = etree.parse(str(project.xml_path))
    assert not tree.xpath(
        ".//p:TextLine[@id='line-1']",
        namespaces={"p": PAGE_2013_NAMESPACE},
    )


def test_page_auto_correct_refuses_dirty_document(qtbot, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    project = _project(tmp_path)
    controller = _open_controller(qtbot, project)
    document = controller.session.document
    assert document is not None
    document.line_by_id("line-1").text = "edited"
    errors: list[str] = []
    controller._error = lambda title, message: errors.append(title)  # type: ignore[method-assign]
    controller.auto_correct_page()
    assert errors == ["Save or discard current changes"]
    assert controller._current_run() is None


def test_batch_result_does_not_replace_xml_saved_while_worker_ran(
    qtbot, tmp_path: Path
) -> None:  # type: ignore[no-untyped-def]
    project = _project(tmp_path)
    controller = _open_controller(qtbot, project)
    live = controller.session.document
    pair = controller.session.current_pair
    assert live is not None and pair is not None
    assert controller.ground_truth is not None

    worker_document = parse_page(project.xml_path)
    worker_document.image_path = pair.image_path
    proposal = controller.workflow.propose(worker_document, controller.ground_truth)
    source_digest = _source_digest(project.xml_path)

    project.xml_path.write_bytes(_page_xml("saved while batch ran"))
    controller._apply_batch_proposals(
        _BatchResult(
            (_BatchProposal(pair, worker_document, proposal, source_digest),)
        )
    )

    assert controller.session.document is live
    assert controller.documents[project.xml_path] is live
    assert project.xml_path not in controller.runs
    assert _saved_unicode(project.xml_path) == "saved while batch ran"


def test_page_review_ignores_report_only_matches(
    qtbot, tmp_path: Path
) -> None:  # type: ignore[no-untyped-def]
    project = _project(tmp_path)
    ground_truth = Document()
    ground_truth.add_paragraph("[1r]")
    ground_truth.add_paragraph("قديم")
    assert project.paths.ground_truth_path is not None
    ground_truth.save(project.paths.ground_truth_path)
    controller = _open_controller(qtbot, project)
    _complete_current_page_correction(controller)
    before = controller.window.undo_stack.count()

    controller.keep_line("line-1")
    controller.reject_line("line-1")
    controller.keep_page()
    controller.reject_page()

    assert controller.window.undo_stack.count() == before
    assert controller.window.undo_stack.isClean()
