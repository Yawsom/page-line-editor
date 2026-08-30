from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from page_line_editor.application.auto_workflow import (
    AutoCorrectionWorkflow,
    ReviewDecision,
    StaleCorrectionProposal,
)
from page_line_editor.application.ground_truth import parse_ground_truth_docx
from page_line_editor.pagexml.parser import parse_page


def page_xml(lines: str, *, image: str = "1r.jpg") -> str:
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<PcGts xmlns="http://schema.primaresearch.org/PAGE/gts/pagecontent/2013-07-15">
  <Metadata><Creator>test</Creator><Created>2025-01-01T00:00:00Z</Created><LastChange>2025-01-01T00:00:00Z</LastChange></Metadata>
  <Page imageFilename="{image}" imageWidth="1000" imageHeight="1200">
    <TextRegion id="r1">
      <Coords points="0,0 900,0 900,1000 0,1000"/>
      {lines}
    </TextRegion>
  </Page>
</PcGts>'''


def text_line(
    line_id: str,
    text: str,
    *,
    coords: str = "100,100 700,100 700,160 100,160",
    baseline: str = "120,145 680,145",
) -> str:
    return f'''<TextLine id="{line_id}">
      <Coords points="{coords}"/>
      <Baseline points="{baseline}"/>
      <TextEquiv><Unicode>{text}</Unicode></TextEquiv>
    </TextLine>'''


def write_xml(path: Path, lines: str, *, image: str = "1r.jpg") -> bytes:
    content = page_xml(lines, image=image).encode()
    path.write_bytes(content)
    return content


def write_ground_truth(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def test_ground_truth_docx_parser_and_page_resolution(tmp_path: Path) -> None:
    source = tmp_path / "truth.docx"
    write_ground_truth(
        source,
        ["ignored preface", "[1R]", " السطر الأول ", "", "[2v]", "السطر الثاني"],
    )
    book = parse_ground_truth_docx(source)
    xml_path = tmp_path / "unrelated.xml"
    write_xml(xml_path, text_line("l1", "السطر القديم"), image="1r.jpg")
    document = parse_page(xml_path)

    assert tuple(book.pages) == ("1r", "2v")
    assert [line.text for line in book.lines_for_document(document)] == ["السطر الأول"]
    assert book.pages["1r"][0].index == 0


def test_ground_truth_strips_bidi_reads_tables_and_warns_on_folio_like(
    tmp_path: Path,
) -> None:
    source = tmp_path / "truth.docx"
    document = Document()
    document.add_paragraph("\u200f[93V]\u200e")
    document.add_paragraph("سطر الصفحة")
    document.add_paragraph("[ 99v ]")
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "[2r]"
    table.cell(1, 0).text = "سطر الجدول"
    document.save(source)

    book = parse_ground_truth_docx(source)
    assert tuple(book.pages) == ("93v", "2r")
    assert [line.text for line in book.pages["93v"]] == ["سطر الصفحة"]
    assert [line.text for line in book.pages["2r"]] == ["سطر الجدول"]
    assert any("99v" in warning for warning in book.warnings)


def test_run_applies_in_memory_audits_and_never_overwrites_xml(tmp_path: Path) -> None:
    xml_path = tmp_path / "1r.xml"
    original_bytes = write_xml(xml_path, text_line("line-7", "النسخة القديمة"))
    truth_path = tmp_path / "truth.docx"
    write_ground_truth(truth_path, ["[1r]", "النسخة الجديده"])
    document = parse_page(xml_path)
    line = document.line_by_id("line-7")
    line.proposal_id = "prior-proposal"
    line.proposal_state = "prior-state"
    line.diff_text = "prior diff"
    line.pre_correction_text = "prior text"

    workflow = AutoCorrectionWorkflow()
    run = workflow.run_page(
        document,
        workflow.load_ground_truth(truth_path),
        tmp_path / "history",
    )

    assert line.text == "النسخة الجديده"
    assert line.proposal_state == "applied"
    assert line.pre_correction_text == "النسخة القديمة"
    assert run.application_for_line("line-7").decision is ReviewDecision.APPLIED
    assert xml_path.read_bytes() == original_bytes
    assert run.audit.original_xml.read_bytes() == original_bytes
    assert run.audit.original_xml.parent.name == "originals"
    assert run.audit.json_report.exists()
    assert run.audit.html_index.exists()
    payload = json.loads(run.audit.json_report.read_text(encoding="utf-8"))
    assert "1r.xml::line-7" in payload["records"]
    manifest = json.loads(run.audit.manifest.read_text(encoding="utf-8"))
    assert manifest["records"]["1r.xml::line-7"]["decision"] == "applied"

    run.reject_line("line-7")
    assert line.text == "النسخة القديمة"
    assert line.proposal_id == "prior-proposal"
    assert line.proposal_state == "prior-state"
    assert line.diff_text == "prior diff"
    assert line.pre_correction_text == "prior text"
    assert xml_path.read_bytes() == original_bytes


def test_extra_line_is_removed_from_active_page_only_after_keep(
    tmp_path: Path,
) -> None:
    xml_path = tmp_path / "1r.xml"
    write_xml(
        xml_path,
        text_line(
            "noise",
            "1",
            coords="100,100 120,100 120,130 100,130",
            baseline="102,125 118,125",
        ),
    )
    document = parse_page(xml_path)
    run = AutoCorrectionWorkflow().run_page(document, (), tmp_path / "history")
    line = document.line_by_id("noise")

    assert run.application_for_line("noise").decision is ReviewDecision.PENDING
    assert not line.deleted
    assert line.correction_status == "EXTRA"
    assert line.proposal_state == "pending"
    assert len(document.active_lines) == 1

    run.keep_line("noise")

    assert run.application_for_line("noise").decision is ReviewDecision.KEPT
    assert line.deleted
    assert line.correction_status == "REMOVED"
    assert document.active_lines == []


def test_rejecting_pending_extra_preserves_line(tmp_path: Path) -> None:
    xml_path = tmp_path / "1r.xml"
    write_xml(xml_path, text_line("extra", "هامش زائد"))
    document = parse_page(xml_path)
    run = AutoCorrectionWorkflow().run_page(document, (), tmp_path / "history")

    run.reject_line("extra")

    line = document.line_by_id("extra")
    assert run.application_for_line("extra").decision is ReviewDecision.REJECTED
    assert not line.deleted
    assert line.proposal_state == ""
    assert document.active_lines == [line]


def test_keep_line_and_page_confirm_automatically_applied_values(tmp_path: Path) -> None:
    xml_path = tmp_path / "1r.xml"
    original_bytes = write_xml(
        xml_path,
        text_line("l1", "السطر القديم", coords="100,100 700,100 700,160 100,160")
        + text_line(
            "l2",
            "النص السابق",
            coords="100,250 700,250 700,310 100,310",
            baseline="120,295 680,295",
        ),
    )
    document = parse_page(xml_path)
    run = AutoCorrectionWorkflow().run_page(
        document,
        ("السطر الجديد", "النص اللاحق"),
        tmp_path / "history",
    )

    run.keep_line("l1")
    assert run.application_for_line("l1").decision is ReviewDecision.KEPT
    assert document.line_by_id("l1").text == "السطر الجديد"
    run.keep_page()
    assert all(
        application.decision in {ReviewDecision.KEPT, ReviewDecision.REPORT_ONLY}
        for application in run.applications
    )
    assert document.line_by_id("l2").text == "النص اللاحق"
    assert xml_path.read_bytes() == original_bytes


def test_rejecting_either_merge_member_restores_all_line_states(tmp_path: Path) -> None:
    xml_path = tmp_path / "1r.xml"
    original_bytes = write_xml(
        xml_path,
        text_line(
            "l1",
            "ورحمة الله",
            coords="100,100 380,100 380,160 100,160",
            baseline="110,145 240,142 370,145",
        )
        + text_line(
            "l2",
            "السلام عليكم",
            coords="500,101 800,101 800,161 500,161",
            baseline="510,146 650,141 790,146",
        ),
    )
    document = parse_page(xml_path)
    before = {
        line.id: (line.text, line.polygon, line.baseline, line.deleted)
        for line in document.lines
    }
    run = AutoCorrectionWorkflow().run_page(
        document,
        ("السلام عليكم ورحمة الله",),
        tmp_path / "history",
    )

    application = run.application_for_line("l1")
    assert application.proposal.status.value == "MERGE"
    assert any(line.deleted for line in document.lines)
    run.reject_line("l1")

    assert {
        line.id: (line.text, line.polygon, line.baseline, line.deleted)
        for line in document.lines
    } == before
    assert application.decision is ReviewDecision.REJECTED
    assert xml_path.read_bytes() == original_bytes


def test_reject_page_restores_every_applied_line(tmp_path: Path) -> None:
    xml_path = tmp_path / "1r.xml"
    original_bytes = write_xml(
        xml_path,
        text_line("l1", "السطر القديم")
        + text_line(
            "l2",
            "النص السابق",
            coords="100,250 700,250 700,310 100,310",
            baseline="120,295 680,295",
        ),
    )
    document = parse_page(xml_path)
    original_text = {line.id: line.text for line in document.lines}
    run = AutoCorrectionWorkflow().run_page(
        document,
        ("السطر الجديد", "النص اللاحق"),
        tmp_path / "history",
    )

    run.reject_page()

    assert {line.id: line.text for line in document.lines} == original_text
    assert all(
        application.decision in {ReviewDecision.REJECTED, ReviewDecision.REPORT_ONLY}
        for application in run.applications
    )
    assert xml_path.read_bytes() == original_bytes


def test_stale_proposal_is_rejected_before_audit_or_mutation(tmp_path: Path) -> None:
    xml_path = tmp_path / "1r.xml"
    original_bytes = write_xml(xml_path, text_line("l1", "السطر القديم"))
    document = parse_page(xml_path)
    workflow = AutoCorrectionWorkflow()
    proposal = workflow.propose(document, ("السطر الجديد",))
    document.line_by_id("l1").text = "تعديل يدوي لاحق"

    with pytest.raises(StaleCorrectionProposal):
        workflow.apply(document, proposal, tmp_path / "history")

    assert document.line_by_id("l1").text == "تعديل يدوي لاحق"
    assert not (tmp_path / "history").exists()
    assert xml_path.read_bytes() == original_bytes


def test_apply_nfc_normalizes_ground_truth_text(tmp_path: Path) -> None:
    import unicodedata

    composed = "café"
    decomposed = unicodedata.normalize("NFD", composed)
    assert decomposed != composed
    xml_path = tmp_path / "1r.xml"
    write_xml(xml_path, text_line("l1", "cafe"))
    truth_path = tmp_path / "truth.docx"
    write_ground_truth(truth_path, ["[1r]", decomposed])
    document = parse_page(xml_path)
    AutoCorrectionWorkflow().run_page(
        document,
        parse_ground_truth_docx(truth_path),
        tmp_path / "history",
    )
    assert document.line_by_id("l1").text == composed
